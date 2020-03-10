from docx import Document
import os
import docxToPdf as createPdf
from datetime import datetime

# Variables
COMPANY_NAME = "Procom"
TIME = datetime.now().strftime("%b %Y %d")
V = "Vancouver, BC"
R = "Richmond, BC"
B = "Burnaby, BC"
VI = "Victoria, BC"

DICT = {
    "send_date": TIME,
    "namecompany": COMPANY_NAME,
    "company_address": R
}


OLDPATH = "C:\\Users\\horatio xu\\Desktop\\files\\Career\\CV"
NEW_FOLDER_PATH = "C:\\Users\horatio xu\\Desktop\\files\\Career\\Company\\" + COMPANY_NAME


def main():
    mkdir(NEW_FOLDER_PATH)
    file_path = NEW_FOLDER_PATH
    for fileName in os.listdir(OLDPATH):
        oldFile = OLDPATH + "\\" + fileName
        newFile = file_path + "\\" + fileName
        if oldFile.split(".")[1] == 'docx':
            document = Document(oldFile)
            document = check(document)
            document.save(newFile)
            createPdf.docx_to_pdf(newFile, file_path + "\\" + "cover_letter.pdf")


def check(document):
    # tables
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                for key, value in DICT.items():
                    if key in table.cell(row, col).text:
                        print(key + "->" + value)
                        table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

    # paragraphs
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in DICT.items():
                if key in para.runs[i].text:
                    print(key + "->" + value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    return document


# Create a new folder
def mkdir(path):
    folder = os.path.exists(path)
    if not folder:
        os.makedirs(path)
        print('folder created')

    else:
        print(path + 'folder already exited')


if __name__ == '__main__':
    main()
